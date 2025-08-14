from io import BytesIO
import logging

from docx import Document
from docx.shared import Inches
import pypandoc

from ..infrastructure.repositories.analysis_repository import AnalysisRepository
from ..api.schemas import AnalysisExportDTO, AnalysisStepExportDTO
from .blob_storage_service import BlobStorageService
from .analysis_service import AnalysisNotFoundException


class ExportService:
    def __init__(
        self,
        analysis_repo: AnalysisRepository,
        blob_storage_service: BlobStorageService,
    ) -> None:
        self.analysis_repo = analysis_repo
        self.blob_storage_service = blob_storage_service

    async def generate_word_document(
        self, analysis_detail: AnalysisExportDTO, content_type: str = "assembly"
    ) -> BytesIO:
        """
        Génère un document Word à partir des détails d'une analyse.

        Args:
            analysis_detail: Détails de l'analyse
            content_type: Type de contenu à inclure ('transcription' ou 'assembly')

        Returns:
            BytesIO: Buffer contenant le document Word
        """
        import tempfile
        import os
        
        # Initialiser une liste pour accumuler les parties du document en Markdown
        markdown_parts = []
        
        # Titre du document
        markdown_parts.append(f"# Analyse: {analysis_detail.filename}\n\n")

        if content_type == "transcription":
            # Inclure uniquement le titre et la transcription
            markdown_parts.append("# Transcription\n\n")
            if analysis_detail.transcript:
                # Échapper les caractères spéciaux qui pourraient poser problème
                transcript = analysis_detail.transcript.replace('\\', '\\\\')
                markdown_parts.append(transcript)
            else:
                markdown_parts.append("Le contenu de la transcription n'a pas pu être récupéré.")
        else:  # content_type == "assembly"
            # Utiliser le nom du prompt s'il est disponible, sinon un titre par défaut
            prompt_title = analysis_detail.prompt_name if analysis_detail.prompt_name else "Assemblage des résultats"
            markdown_parts.append(f"# {prompt_title}\n\n")

            # Itérez sur la liste analysis_detail.steps
            for step in analysis_detail.steps:
                # Ajoutez le nom de l'étape comme un titre de niveau 2
                markdown_parts.append(f"## {step.step_name}\n\n")
                # Ajoutez le contenu de l'étape
                if step.content:
                    # Échapper les caractères spéciaux qui pourraient poser problème
                    content = step.content.replace('\\', '\\\\')
                    markdown_parts.append(content)
                markdown_parts.append("\n\n")

        # Joindre toutes les parties en une seule chaîne Markdown
        full_markdown_text = "".join(markdown_parts)
        
        # Convertir le Markdown en document Word
        buffer = BytesIO()
        try:
            # Créer un fichier temporaire pour la conversion
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                tmp_file_path = tmp_file.name
            
            # Convertir le markdown en docx en utilisant le fichier temporaire
            pypandoc.convert_text(
                full_markdown_text, format='markdown', to='docx', outputfile=tmp_file_path
            )
            
            # Lire le contenu du fichier temporaire dans le buffer
            with open(tmp_file_path, 'rb') as f:
                buffer.write(f.read())
            
            # Supprimer le fichier temporaire
            os.unlink(tmp_file_path)
        except FileNotFoundError:
            logging.error("Pandoc non trouvé. Impossible de convertir le Markdown.")
            # Créez un document de fallback simple
            fallback_doc = Document()
            fallback_doc.add_paragraph("Erreur: Pandoc n'est pas installé sur le serveur.")
            fallback_doc.add_paragraph(full_markdown_text)
            fallback_doc.save(buffer)
        except Exception as e:
            logging.error(f"Erreur lors de la conversion du Markdown: {str(e)}")
            # Créez un document de fallback simple avec le contenu brut
            fallback_doc = Document()
            fallback_doc.add_paragraph("Erreur lors de la conversion du document.")
            fallback_doc.add_paragraph("Contenu brut:")
            fallback_doc.add_paragraph(full_markdown_text)
            fallback_doc.save(buffer)
            
        buffer.seek(0)
        return buffer

    async def get_analysis_detail_for_export(
        self, analysis_id: str, user_id: int
    ) -> AnalysisExportDTO:
        """
        Récupère les détails d'une analyse pour l'export.
        """
        # Conservez la récupération de l'analyse et les vérifications de permissions
        analysis = await self.analysis_repo.get_detailed_by_id(analysis_id)
        if not analysis:
            raise AnalysisNotFoundException("Analysis not found")
        if analysis.user_id != user_id:
            raise PermissionError("Access denied")

        # Récupérer le contenu de la transcription depuis le blob storage
        transcript_content = ""
        if analysis.transcript_blob_name:
            try:
                transcript_content = await self.blob_storage_service.download_blob_as_text(
                    analysis.transcript_blob_name
                )
            except Exception:
                # En cas d'erreur de téléchargement, transcript_content reste une chaîne vide
                pass

        # Conservez la logique qui trie les versions de l'analyse pour trouver la plus récente
        versions_sorted = sorted(
            analysis.versions or [], key=lambda v: v.created_at or 0, reverse=True
        )

        # Initialisez une liste vide
        steps_for_export = []

        # Si versions_sorted n'est pas vide, prenez la première version
        if versions_sorted:
            latest_version = versions_sorted[0]

            # Itérez sur les steps de latest_version
            for step in latest_version.steps or []:
                # Créez une instance de AnalysisStepExportDTO avec step_name et content
                if step.content:  # Only add steps that have content
                    step_dto = AnalysisStepExportDTO(
                        step_name=step.step_name, content=step.content
                    )
                    # Ajoutez-la à steps_for_export
                    steps_for_export.append(step_dto)

        # Récupérer le nom du prompt utilisé
        prompt_name = None
        if versions_sorted and versions_sorted[0].prompt_used:
            prompt_name = versions_sorted[0].prompt_used

        # À la fin, retournez une instance de AnalysisExportDTO
        return AnalysisExportDTO(
            id=analysis.id,
            filename=analysis.filename,
            status=analysis.status,
            created_at=analysis.created_at,
            transcript=transcript_content,
            steps=steps_for_export,
            prompt_name=prompt_name,
        )

    
